
import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import io
import base64
import os
from docx import Document
from docx.shared import Inches

# Set page config
st.set_page_config(
    page_title="Business Contact Dashboard",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Load data and translations
@st.cache_data
def load_data_and_translations():
    df = pd.read_csv("parsed_data.csv")
    translations = {}
    try:
        with open("category_translations.txt", "r", encoding="utf-8") as f:
            for line in f:
                parts = line.strip().split("\\")
                if len(parts) == 2:
                    translations[parts[0].strip()] = parts[1].strip()
    except FileNotFoundError:
        st.error("category_translations.txt not found. Please ensure data parsing is complete.")
    return df, translations

# Initialize session state for data
if 'df' not in st.session_state:
    st.session_state.df, st.session_state.translations = load_data_and_translations()

# Analytics functions
def perform_analytics(df):
    companies_per_category = df["Category"].value_counts().reset_index()
    companies_per_category.columns = ["Category", "Company Count"]
    
    most_populated_category = companies_per_category.iloc[0]
    least_populated_category = companies_per_category.iloc[-1]
    total_companies = df.shape[0]
    total_categories = df["Category"].nunique()
    
    return {
        "companies_per_category": companies_per_category,
        "most_populated_category": most_populated_category,
        "least_populated_category": least_populated_category,
        "total_companies": total_companies,
        "total_categories": total_categories,
    }

# CRUD Functions
def add_company(category, company_name):
    new_row = pd.DataFrame({"Category": [category], "Company": [company_name]})
    st.session_state.df = pd.concat([st.session_state.df, new_row], ignore_index=True)
    st.success(f"Added '{company_name}' to '{category}' category!")

def delete_company(category, company_name):
    mask = (st.session_state.df["Category"] == category) & (st.session_state.df["Company"] == company_name)
    if mask.any():
        st.session_state.df = st.session_state.df[~mask]
        st.success(f"Deleted '{company_name}' from '{category}' category!")
    else:
        st.error(f"Company '{company_name}' not found in '{category}' category!")

def add_category(category_name):
    if category_name not in st.session_state.df["Category"].values:
        # Add a placeholder entry for the new category
        new_row = pd.DataFrame({"Category": [category_name], "Company": ["[New Category - Add Companies]"]})
        st.session_state.df = pd.concat([st.session_state.df, new_row], ignore_index=True)
        st.success(f"Added new category '{category_name}'!")
    else:
        st.warning(f"Category '{category_name}' already exists!")

def delete_category(category_name):
    mask = st.session_state.df["Category"] == category_name
    if mask.any():
        st.session_state.df = st.session_state.df[~mask]
        st.success(f"Deleted category '{category_name}' and all its companies!")
    else:
        st.error(f"Category '{category_name}' not found!")

def export_to_word():
    doc = Document()
    doc.add_heading("Business Contact List", 0)
    
    categories = st.session_state.df["Category"].unique()
    for category in categories:
        doc.add_heading(category, level=1)
        companies = st.session_state.df[st.session_state.df["Category"] == category]["Company"].values
        for company in companies:
            p = doc.add_paragraph()
            p.add_run(f"• {company}")
    
    doc.save("updated_contact_list.docx")
    return "updated_contact_list.docx"

def export_to_excel():
    filename = "updated_contact_list.xlsx"
    st.session_state.df.to_excel(filename, index=False)
    return filename

# Main app
def main():
    st.title("📊 Business Contact Dashboard")
    st.markdown("---")
    
    # Get current data and analytics
    df = st.session_state.df
    translations = st.session_state.translations
    analytics = perform_analytics(df)
    
    # KPIs at the top
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Categories", analytics["total_categories"])
    
    with col2:
        st.metric("Total Companies", analytics["total_companies"])
    
    with col3:
        st.metric("Most Populated Category", 
                 analytics["most_populated_category"]["Category"],
                 f"{analytics['most_populated_category']['Company Count']} companies")
    
    with col4:
        st.metric("Least Populated Category", 
                 analytics["least_populated_category"]["Category"],
                 f"{analytics['least_populated_category']['Company Count']} companies")
    
    st.markdown("---")
    
    # Sidebar for filters and CRUD operations
    st.sidebar.header("🔍 Filters & Search")
    
    # Category filter
    categories = ["All"] + list(df["Category"].unique())
    selected_category = st.sidebar.selectbox("Select Category", categories)
    
    # Company search
    search_term = st.sidebar.text_input("Search Company Name")
    
    # Filter data
    filtered_df = df.copy()
    if selected_category != "All":
        filtered_df = filtered_df[filtered_df["Category"] == selected_category]
    
    if search_term:
        filtered_df = filtered_df[filtered_df["Company"].str.contains(search_term, case=False, na=False)]
    
    # CRUD Operations in Sidebar
    st.sidebar.markdown("---")
    st.sidebar.header("✏️ CRUD Operations")
    
    # Add Company
    with st.sidebar.expander("➕ Add Company"):
        add_category_select = st.selectbox("Select Category for New Company", df["Category"].unique(), key="add_cat")
        add_company_name = st.text_input("Company Name", key="add_comp")
        if st.button("Add Company"):
            if add_company_name:
                add_company(add_category_select, add_company_name)
                st.rerun()
    
    # Delete Company
    with st.sidebar.expander("🗑️ Delete Company"):
        del_category_select = st.selectbox("Select Category", df["Category"].unique(), key="del_cat")
        companies_in_category = df[df["Category"] == del_category_select]["Company"].unique()
        del_company_select = st.selectbox("Select Company to Delete", companies_in_category, key="del_comp")
        if st.button("Delete Company"):
            delete_company(del_category_select, del_company_select)
            st.rerun()
    
    # Add Category
    with st.sidebar.expander("📁 Add Category"):
        new_category_name = st.text_input("New Category Name", key="new_cat")
        if st.button("Add Category"):
            if new_category_name:
                add_category(new_category_name)
                st.rerun()
    
    # Delete Category
    with st.sidebar.expander("🗂️ Delete Category"):
        del_category_name = st.selectbox("Select Category to Delete", df["Category"].unique(), key="del_cat_name")
        if st.button("Delete Category", type="secondary"):
            delete_category(del_category_name)
            st.rerun()
    
    # Export Options
    st.sidebar.markdown("---")
    st.sidebar.header("📤 Export Options")
    
    col1, col2 = st.sidebar.columns(2)
    with col1:
        if st.button("Export to Word"):
            word_file = export_to_word()
            with open(word_file, "rb") as file:
                st.download_button(
                    label="📄 Download Word",
                    data=file.read(),
                    file_name=word_file,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    with col2:
        if st.button("Export to Excel"):
            excel_file = export_to_excel()
            with open(excel_file, "rb") as file:
                st.download_button(
                    label="📊 Download Excel",
                    data=file.read(),
                    file_name=excel_file,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
    
    # Main content area
    tab1, tab2, tab3 = st.tabs(["📈 Analytics", "📋 Data View", "☁️ Word Cloud"])
    
    with tab1:
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Companies by Category (Bar Chart)")
            fig_bar = px.bar(
                analytics["companies_per_category"], 
                x="Category", 
                y="Company Count",
                title="Number of Companies per Category",
                color="Company Count",
                color_continuous_scale="viridis"
            )
            fig_bar.update_layout(xaxis_tickangle=-45)
            st.plotly_chart(fig_bar, use_container_width=True)
        
        with col2:
            st.subheader("Companies by Category (Donut Chart)")
            fig_donut = px.pie(
                analytics["companies_per_category"], 
                values="Company Count", 
                names="Category",
                title="Distribution of Companies by Category",
                hole=0.4
            )
            st.plotly_chart(fig_donut, use_container_width=True)
    
    with tab2:
        st.subheader("📋 Company Data")
        st.write(f"Showing {len(filtered_df)} of {len(df)} companies")
        
        # Display filtered data
        st.dataframe(filtered_df, use_container_width=True)
        
        # Download button
        csv = filtered_df.to_csv(index=False)
        st.download_button(
            label="📥 Download filtered data as CSV",
            data=csv,
            file_name="filtered_companies.csv",
            mime="text/csv"
        )
    
    with tab3:
        st.subheader("☁️ Company Names Word Cloud")
        if len(filtered_df) > 0:
            # Generate word cloud
            text = " ".join(filtered_df["Company"].values)
            wordcloud = WordCloud(width=800, height=400, background_color="white").generate(text)
            
            # Display word cloud
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.imshow(wordcloud, interpolation="bilinear")
            ax.axis("off")
            st.pyplot(fig)
        else:
            st.info("No companies found for the selected filters.")

    st.markdown("---")
    st.subheader("Category Translations")
    for eng, arab in translations.items():
        st.write(f"**{eng}**: {arab}")

if __name__ == "__main__":
    main()